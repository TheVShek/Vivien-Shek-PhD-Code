from tkinter import *
import requests
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX


def getExons(id_, species):
    server = "https://parasite.wormbase.org"

    # species = "Heligmosomoides_polygyrus_"+species
    # id_ = "HPOL_0000515001"

    ext = "/rest-16/lookup/id/" + id_ + "?expand=1"
    print(server + ext)
    r = requests.get(
        server + ext,
        headers={
            "Content-Type": "application/json",
        },
    )

    decoded = r.json()
    print(decoded)
    species = decoded["species"]
    exons = decoded["Transcript"][0]["Exon"]
    exons_ = []
    for e in exons:
        region = (
            "/"
            + e["seq_region_name"]
            + ":"
            + str(e["start"])
            + "-"
            + str(e["end"])
            + ":"
            + str(e["strand"])
            + "?"
        )
        ext = "/rest-16/sequence/region/" + species + region
        print(server + ext)
        r = requests.get(
            server + ext, headers={"Content-Type": "text/plain", "Accept": ""}
        )
        exons_.append(r.text)
    return exons_


def getTranslation(dna):
    r = requests.post(
        "https://web.expasy.org/cgi-bin/translate/dna2aa.cgi",
        data={"dna_sequence": dna, "output_format": "fasta"},
    )
    seqs = r.text.split(">")[1:]
    max_ = 0
    idx = 0
    for i in range(len(seqs)):
        s = "".join(seqs[i].split("\n")[1:])
        tmpMax = re.findall("M[^-]*", s)
        if len(tmpMax) > 0:
            tmpMax = len(max(tmpMax, key=len))
        else:
            continue

        if tmpMax > max_:
            idx = i
            max_ = tmpMax
    return seqs[idx]


class ExonAnnotate:
    def __init__(self, exons, dna):
        self.spans = []
        self.col = WD_COLOR_INDEX.YELLOW
        self.currentSpan = None
        print(dna.upper())
        for e in exons:
            m = re.search(e, dna.upper())
            if m:
                print(e)
                self.spans.append(re.search(e, dna.upper()).span())
        for e in self.spans:
            print(e)

    def color(self):
        if self.col == "yellow":
            return RGBColor(0x00, 0x00, 0xFF)
        else:
            return RGBColor(0xFF, 0x00, 0x00)

    def annotate(self, idx):
        for span in self.spans:
            if idx >= span[0] and idx < span[1]:
                if self.currentSpan == None:
                    self.currentSpan = span
                elif span[0] != self.currentSpan[0]:
                    self.currentSpan = span
                    if self.col == "yellow":
                        self.col = "red"
                    else:
                        self.col = "yellow"
                return True

        return False


class CCPAnnotate:
    def __init__(self, transation, ccps):

        self.annotateRegion = []
        print(transation)
        for ccp in ccps:
            print(ccp)
            m = re.search(ccp, transation)
            self.annotateRegion.append(m.span())

    def annotate(self, idx):
        for a in self.annotateRegion:
            if idx >= a[0] and idx < a[1]:
                return True
        return False


def doOne(f, document):
    f = f.split(">")[1:]
    print(f[0].split("\n")[0])
    # species = re.search("prjeb[^_|]*",f[0].split("\n")[0]).group()
    # print (species)
    species = "prjeb15396"
    id_ = re.search("(HPOL|HPBE)_[1234567890]*", f[0].split("\n")[0]).group()
    print(id_)
    dna = "".join(f[0].split("\n")[1:])

    transation = getTranslation(dna)

    exons = getExons(id_, species)
    print(dna)
    print(exons)
    ccps = ["".join(x.split("\n")[1:]).strip() for x in f[1:]]

    frameStart = int(
        re.search("[123]", re.search("Frame [123]", transation).group()).group()
    )
    print(frameStart)

    exonAnnotate = ExonAnnotate(exons, dna)
    ccpAnnotate = CCPAnnotate("".join(transation.split("\n")[1:]), ccps)

    para = document.add_paragraph()

    transation = "".join(transation.split("\n")[1:])

    for idx in range(57 + frameStart - 1):
        r = para.add_run(dna[idx].lower())
        if exonAnnotate.annotate(idx):
            r.font.color.rgb = exonAnnotate.color()
            r.font.underline = True

    para = document.add_paragraph(" " * (frameStart - 1))

    idx = 0
    while idx * 3 < 57:
        r = para.add_run(" " + transation[idx] + " ")
        if ccpAnnotate.annotate(idx):
            r.font.highlight_color = WD_COLOR_INDEX.GRAY_50
        idx = idx + 1

    # paragraph = document.add_paragraph(translateSpaces[0:57+frameStart-1])

    start = 57 + frameStart - 1
    end = start + 57

    while start < len(dna):
        para = document.add_paragraph()
        for idx in range(start, end):
            r = para.add_run(dna[idx].lower())
            if exonAnnotate.annotate(idx):
                r.font.color.rgb = exonAnnotate.color()
                r.font.underline = True
        para = document.add_paragraph()
        idx = int(start / 3)
        while idx * 3 < end - 2:
            try:
                r = para.add_run(" " + transation[idx] + " ")
            except IndexError:
                break
            if ccpAnnotate.annotate(idx):
                r.font.highlight_color = WD_COLOR_INDEX.GRAY_50
            if transation[idx] == "C" or transation[idx] == "W":
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.bold = True
                r.font.highlight_color = WD_COLOR_INDEX.BLACK
            idx = idx + 1
        # paragraph = document.add_paragraph(translateSpaces[start:end])
        start = end
        if len(dna) < start + 57:
            end = len(dna)
        else:
            end = start + 57


o = open("List of VS CCP HITS SEQ DNA.txt")
o = [">" + x for x in o.read().split(">")[1:]]

document = Document()
style = document.styles["Normal"]
style.font.name = "Courier New"
style.font.size = Pt(12)
style.paragraph_format.space_before = Pt(3)
style.paragraph_format.space_after = Pt(3)

for f in o:
    if "HPOL" not in f and "HPBE" not in f:
        continue
    document.add_paragraph(f.split("\n")[0])
    doOne(f, document)

document.save("CCPExons.docx")
