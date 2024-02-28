from docx import Document
from docx.shared import RGBColor
import regex as re
from docx.enum.text import WD_COLOR_INDEX

red = RGBColor(0xFF, 0x00, 0x00)
green = RGBColor(0x00, 0xFF, 0x00)
grey = RGBColor(195, 195, 195)


def annotateSeq(p, seq, highlight=False):
    for c in seq:
        run = p.add_run(c)
        if c == "C":
            run.font.color.rgb = green
        elif c == "W":
            run.font.color.rgb = red
        if highlight:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def annotateSeqObj(p, seq):
    for i in range(len(seq.seq)):
        c = seq.seq[i]
        run = p.add_run(c)
        if c == "C":
            run.font.color.rgb = green
        elif c == "W":
            run.font.color.rgb = red
        if seq.annotateRegions.annotate(i):
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW


class Sequence:
    def __init__(self, seq):
        self.title = ">" + seq.split("\n")[0]
        self.seq = "".join(seq.split("\n")[1:])
        self.allPossCCPS = re.finditer(
            "C[^C]*C[^C]*C[^CW]*W[^C]*C", self.seq, overlapped=True
        )
        self.annotateRegions = self.getAnnotatable(
            re.finditer("C[^C]*C[^C]*C[^CW]*W[^C]*C", self.seq, overlapped=True)
        )

    def getPossCCPS(self, seq):
        return re.finditer("C[^C]*C[^C]*C[^CW]*W[^C]*C", seq, overlapped=True)

    def getAnnotatable(self, possCCPS):
        shortCCPS = [x for x in possCCPS if len(x.group()) < 115]
        return AnnotateRegion([x for x in shortCCPS])


class AnnotateRegion:
    def __init__(self, shorttCCPS):
        regions = [x.span() for x in shorttCCPS]
        seqs = [x.group() for x in shorttCCPS]
        nonOverlapping = regions[:]
        idx = 0
        while idx < (len(shorttCCPS) - 1):
            if regions[idx][1] > regions[idx + 1][0]:
                nonOverlapping.remove(regions[idx])
                nonOverlapping.remove(regions[idx + 1])
                seqs.remove(shorttCCPS[idx].group())
                seqs.remove(shorttCCPS[idx + 1].group())
                idx = idx + 2
            else:
                idx = idx + 1
        self.nonOverlapping = nonOverlapping
        self.seqs = seqs

    def annotate(self, i):
        for r in self.nonOverlapping:
            if i >= r[0] and i < r[1]:
                return True
        return False


f = open(
    "All_5_IPR_VSCCP_SP_SEQ.txt",
    "r",
)
f = f.read().split(">")[1:]

document = Document()
print("name,ccp")

for i in f:
    seq = Sequence(i)
    p = document.add_paragraph()
    pr = p.add_run(seq.title + "\n")
    pr.font.bold = True
    pr.font.highlight_color = WD_COLOR_INDEX.PINK
    annotateSeqObj(p, seq)
    # mathes = re.finditer("C[^C]*C[^C]*C[^CW]*W[^C]*C",s,overlapped=True)
    p.add_run("\n")
    n = 1

    for m in seq.allPossCCPS:
        p.add_run("\n")

        allC = [c.start(0) for c in re.finditer("C", m.group())]

        r = p.add_run(
            "output "
            + str(n)
            + " at "
            + str(m.span())
            + " len "
            + str(len(m.group()))
            + ", C1-C2 = "
            + str(allC[1] - allC[0] - 1)
            + ", C2-C3 = "
            + str(allC[2] - allC[1] - 1)
            + ", C3-C4 = "
            + str(allC[3] - allC[2] - 1)
            + ":\n"
        )
        r.font.bold = True

        if m.group() in seq.annotateRegions.seqs:
            annotateSeq(p, m.group(), True)
        else:
            annotateSeq(p, m.group())
        n = n + 1
    print(f"{seq.title},{n-1}")


document.save("All_5_IPR_VSCCP_SP_SEQ.docx")
